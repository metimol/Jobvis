"""Multi-format CV Parser Service for PDF, DOCX, and TXT documents."""

import io
import logging
import re

import docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)

MAX_CV_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB limit


class CVParserService:
    """Multi-format CV Parser extracting raw sanitized text from candidate documents."""

    @classmethod
    def sanitize_text(cls, text: str) -> str:
        """Sanitize extracted text by removing control characters and normalizing whitespace."""
        if not text:
            return ""
        # Remove null bytes and non-printable control characters (except newline, tab)
        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Normalize linebreaks and strip leading/trailing whitespace per line
        cleaned = "\n".join([line.strip() for line in cleaned.splitlines() if line.strip()])
        return cleaned

    @classmethod
    def parse_pdf(cls, file_bytes: bytes) -> str:
        """Extract text from PDF document using pypdf."""
        if not file_bytes:
            return ""
        stream = io.BytesIO(file_bytes)
        try:
            reader = PdfReader(stream)
            extracted_pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(page_text)
            text = "\n".join(extracted_pages)
            return cls.sanitize_text(text)
        except Exception as e:
            logger.error("Failed to parse PDF document: %s", e)
            raise ValueError(f"Corrupted or invalid PDF file: {e}") from e

    @classmethod
    def parse_docx(cls, file_bytes: bytes) -> str:
        """Extract text from Word DOCX document using python-docx."""
        if not file_bytes:
            return ""
        stream = io.BytesIO(file_bytes)
        try:
            doc = docx.Document(stream)
            extracted_paragraphs = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    extracted_paragraphs.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            extracted_paragraphs.append(cell.text.strip())
            text = "\n".join(extracted_paragraphs)
            return cls.sanitize_text(text)
        except Exception as e:
            logger.error("Failed to parse DOCX document: %s", e)
            raise ValueError(f"Corrupted or invalid DOCX file: {e}") from e

    @classmethod
    def parse_txt(cls, file_bytes: bytes) -> str:
        """Extract text from plain text document."""
        if not file_bytes:
            return ""
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode("latin-1")
            except Exception as e:
                logger.error("Failed to decode TXT document: %s", e)
                raise ValueError(f"Unable to decode text file: {e}") from e
        return cls.sanitize_text(text)

    @classmethod
    def parse_document(
        cls,
        file_bytes: bytes,
        filename: str,
        max_size_bytes: int = MAX_CV_FILE_SIZE_BYTES,
    ) -> str:
        """Validate format & size, parse document according to extension, and return sanitized text."""
        if len(file_bytes) > max_size_bytes:
            raise ValueError(f"File size {len(file_bytes)} exceeds limit of {max_size_bytes} bytes")

        if not filename or "." not in filename:
            raise ValueError("Unsupported file format: missing file extension")

        ext = filename.lower().rsplit(".", 1)[-1].strip()

        if ext == "pdf":
            return cls.parse_pdf(file_bytes)
        if ext == "docx":
            return cls.parse_docx(file_bytes)
        if ext in ["txt", "text"]:
            return cls.parse_txt(file_bytes)
        raise ValueError(f"Unsupported file format: .{ext}")


# Alias for backward compatibility
CVParser = CVParserService
